"""
Field Sync - Systems Documentation Folder Builder
Run: pip install python-docx   (one time)
Then: python build_fs.py
Output lands on your Desktop.
"""

import os, zipfile, datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DESKTOP = Path.home() / "Desktop"
ROOT = DESKTOP / "Field Sync – Systems Documentation"
ZIP_PATH = DESKTOP / "Field Sync – Systems Documentation.zip"

STRUCTURE = {
    "FS-00 – ADMIN & INDEXES": [
        "FS-00.01 – Table of Contents (Master Index)",
        "FS-00.02 – Document Naming Rules",
        "FS-00.03 – Version Control Rules",
        "FS-00.04 – Glossary of Terms",
        "FS-00.05 – Brand Voice & Writing Standards",
    ],
    "FS-01 – COMPANY SYSTEMS": [
        "FS-01.01 – Company Overview",
        "FS-01.02 – Mission, Vision, Values",
        "FS-01.03 – Organizational Structure",
        "FS-01.04 – Roles & Responsibilities",
        "FS-01.05 – Communication Standards",
        "FS-01.06 – SOP Template (Master)",
    ],
    "FS-02 – DATA GOVERNANCE & INFORMATION MANAGEMENT": [
        "FS-02.01 – Data Governance Overview",
        "FS-02.02 – Data Classification Framework",
        "FS-02.03 – Data Retention & Deletion Policies",
        "FS-02.04 – PII Handling Standards",
        "FS-02.05 – Client Data Boundaries",
        "FS-02.06 – Storage & Encryption Standards",
        "FS-02.07 – Access Control & Permissions",
        "FS-02.08 – Audit Trails & Monitoring",
        "FS-02.09 – Compliance Requirements",
    ],
    "FS-03 – AI ETHICS, SAFETY & RESPONSIBLE USE": [
        "FS-03.01 – Ethical Principles",
        "FS-03.02 – Responsible AI Guidelines",
        "FS-03.03 – Bias Mitigation",
        "FS-03.04 – Transparency Standards",
        "FS-03.05 – Safety Protocols",
        "FS-03.06 – Human-in-the-Loop Requirements",
        "FS-03.07 – Client Ethical Use Policy",
        "FS-03.08 – Incident Reporting & Escalation",
    ],
    "FS-04 – MODEL MANAGEMENT & QUALITY CONTROL": [
        "FS-04.01 – Model Inventory",
        "FS-04.02 – Versioning & Change Logs",
        "FS-04.03 – Testing & Evaluation Procedures",
        "FS-04.04 – Performance Metrics",
        "FS-04.05 – Drift Monitoring",
        "FS-04.06 – Quality Assurance Standards",
        "FS-04.07 – Review & Approval Workflow",
    ],
    "FS-05 – PRODUCT ARCHITECTURE": [
        "FS-05.01 – Call Flow Logic",
        "FS-05.02 – Hours Saved Calculation",
        "FS-05.03 – Job Booking Logic",
        "FS-05.04 – Follow-Up Logic",
        "FS-05.05 – Spam Filtering Logic",
        "FS-05.06 – Emergency Routing Logic",
        "FS-05.07 – Data Model (High-Level)",
        "FS-05.08 – API / Integration Notes",
    ],
    "FS-06 – AUTOMATION & WORKFLOW ENGINEERING": [
        "FS-06.01 – Automation Overview",
        "FS-06.02 – Workflow Blueprints",
        "FS-06.03 – Trigger & Event Logic",
        "FS-06.04 – Routing Rules",
        "FS-06.05 – Exception Handling",
        "FS-06.06 – Optimization & Efficiency Standards",
        "FS-06.07 – Automation QA Procedures",
    ],
    "FS-07 – INTEGRATIONS & SYSTEM CONNECTIONS": [
        "FS-07.01 – Integration Overview",
        "FS-07.02 – CRM Integrations",
        "FS-07.03 – Phone System Integrations",
        "FS-07.04 – Calendar Integrations",
        "FS-07.05 – Email Integrations",
        "FS-07.06 – Webhooks & Event Triggers",
        "FS-07.07 – API Endpoints (High-Level)",
        "FS-07.08 – Zapier / Make Workflows",
        "FS-07.09 – Data Sync Rules & Mapping",
        "FS-07.10 – Authentication & Security",
        "FS-07.11 – Error Handling & Logging",
        "FS-07.12 – Integration Versioning & Deprecation",
    ],
    "FS-08 – FIELD SYNC LEAGUE SYSTEM": [
        "FS-08.01 – League Overview",
        "FS-08.02 – Leaderboard System",
        "FS-08.03 – Draft Picks System",
        "FS-08.04 – Assistant Card System",
        "FS-08.05 – Employee # / Account # Rules",
        "FS-08.06 – Stats Tracking System",
        "FS-08.07 – Quarterly MVP System",
        "FS-08.08 – Free Agency System",
        "FS-08.09 – Hall of Fame System",
        "FS-08.10 – NPC Assistant Roster",
        "FS-08.11 – Real Assistant Roster",
    ],
    "FS-09 – CLIENT OPERATIONS": [
        "FS-09.01 – Onboarding Process",
        "FS-09.02 – Drafting an Assistant",
        "FS-09.03 – Assistant Setup Checklist",
        "FS-09.04 – Client Portal Overview",
        "FS-09.05 – Season Reports",
        "FS-09.06 – Renewal Process",
        "FS-09.07 – Support Escalation Flow",
    ],
    "FS-10 – BRANDING & DESIGN": [
        "FS-10.01 – Logo Files",
        "FS-10.02 – Color Palette",
        "FS-10.03 – Typography",
        "FS-10.04 – Assistant Portrait Style Guide",
        "FS-10.05 – Card Design Templates",
        "FS-10.06 – Badge System",
        "FS-10.07 – Website Style Guide",
    ],
    "FS-11 – WEBSITE CONTENT": [
        "FS-11.01 – Home Page",
        "FS-11.02 – League Page",
        "FS-11.03 – Draft Room Page",
        "FS-11.04 – Leaderboard Page",
        "FS-11.05 – MVP Winners Page",
        "FS-11.06 – Roster Page",
        "FS-11.07 – Free Agency Page",
        "FS-11.08 – Contact Page",
        "FS-11.09 – FAQ Page",
    ],
    "FS-12 – MARKETING & SALES": [
        "FS-12.01 – Pitch Deck",
        "FS-12.02 – Sales Scripts",
        "FS-12.03 – Email Templates",
        "FS-12.04 – Social Media Templates",
        "FS-12.05 – Case Studies",
        "FS-12.06 – Testimonials",
        "FS-12.07 – Launch Plan",
        "FS-12.08 – Quarterly Campaigns",
    ],
    "FS-13 – LEGAL & COMPLIANCE": [
        "FS-13.01 – Terms of Service",
        "FS-13.02 – Privacy Policy",
        "FS-13.03 – Client Agreement",
        "FS-13.04 – Assistant Licensing Terms",
        "FS-13.05 – Data Handling Notes",
    ],
    "FS-14 – PRODUCT OFFERINGS & AI TOOLS": [
        "FS-14.01 – Product Line Overview",
        "FS-14.02 – AI Assistant Types",
        "FS-14.03 – AI Phone Assistant",
        "FS-14.04 – AI Scheduling Assistant",
        "FS-14.05 – AI CRM Assistant",
        "FS-14.06 – AI Follow-Up Assistant",
        "FS-14.07 – AI Lead Qualification Assistant",
        "FS-14.08 – AI Inbox Assistant",
        "FS-14.09 – AI Reporting Assistant",
        "FS-14.10 – Workflow Automations (Trades)",
        "FS-14.11 – Pricing & Packaging",
        "FS-14.12 – Feature Comparison Matrix",
        "FS-14.13 – Upgrade Paths & Add-Ons",
        "FS-14.14 – Product Roadmap",
    ],
    "FS-15 – PROVIDE SECTION": [
        "FS-15.01 – Provided Mockups",
        "FS-15.02 – Provided Drafts",
        "FS-15.03 – Provided Creative Assets",
        "FS-15.04 – Working Documents",
    ],
    "FS-16 – ARCHIVE & SANDBOX": [
        "FS-16.01 – Deprecated Systems",
        "FS-16.02 – Old Versions",
        "FS-16.03 – Experiments",
        "FS-16.04 – Temporary Notes",
    ],
}

NAVY = RGBColor(0x0A, 0x1F, 0x3C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)
TODAY = datetime.date.today().strftime("%B %d, %Y")


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from lxml import etree
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = etree.SubElement(tc_pr, qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def make_docx(folder_path, code, title):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Header banner
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_table.cell(0, 0)
    set_cell_shading(cell, "0A1F3C")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FIELD SYNC")
    run.font.size = Pt(22)
    run.font.color.rgb = WHITE
    run.bold = True
    run2 = p.add_run("\nSystems Documentation")
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph("")

    # Document title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title_p.add_run(f"{code} – {title}")
    run_t.font.size = Pt(16)
    run_t.font.color.rgb = NAVY
    run_t.bold = True

    doc.add_paragraph("")

    # Metadata table
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    labels = ["Document Code", "Status", "Owner", "Version", "Last Updated"]
    values = [code, "Placeholder – Not Started", "Field Sync", "0.1", TODAY]
    for i, (label, value) in enumerate(zip(labels, values)):
        lc = meta.cell(i, 0)
        lc.text = ""
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = NAVY
        set_cell_shading(lc, "E8EDF3")
        vc = meta.cell(i, 1)
        vc.text = ""
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)
        vr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("")

    # Sections
    for heading in ["Purpose", "Scope", "Content"]:
        h = doc.add_heading(heading, level=2)
        for run in h.runs:
            run.font.color.rgb = NAVY
        doc.add_paragraph("[Enter details here]")
        doc.add_paragraph("")

    # Footer
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "Field Sync – Confidential | AI Consulting for the Trades"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY
    fr.italic = True

    # Save
    safe_name = (
        f"{code} – {title}"
        .replace("/", "-")
        .replace("\\", "-")
        .replace(":", "-")
        .replace("#", "")
    )
    file_path = folder_path / f"{safe_name}.docx"
    doc.save(str(file_path))
    return file_path


def main():
    print("=" * 55)
    print("  Field Sync – Systems Documentation Builder")
    print("=" * 55)

    if ROOT.exists():
        import shutil
        shutil.rmtree(ROOT)

    file_count = 0
    for folder_name, files in STRUCTURE.items():
        folder_path = ROOT / folder_name
        folder_path.mkdir(parents=True, exist_ok=True)
        for entry in files:
            code = entry.split(" – ")[0].strip()
            title = " – ".join(entry.split(" – ")[1:]).strip()
            make_docx(folder_path, code, title)
            file_count += 1
            print(f"  [OK] {code} – {title}")

    print(f"\n  Created {file_count} files in {len(STRUCTURE)} folders.")
    print(f"  Folder: {ROOT}\n")

    # ZIP
    print("  Packaging ZIP...")
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for fpath in sorted(ROOT.rglob("*")):
            zf.write(fpath, fpath.relative_to(DESKTOP))
    print(f"  ZIP saved: {ZIP_PATH}")
    print("\n  Done! Check your Desktop.")
    print("=" * 55)


if __name__ == "__main__":
    main()
